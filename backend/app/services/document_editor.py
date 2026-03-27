import io
import re
from typing import List
from docx import Document
from docx.shared import Pt
from app.models.edit_schemas import TextModification

class DocumentEditor:
    """
    Service class responsible for applying edits to contract text and 
    generating downloadable files.
    """

    def apply_modifications(self, original_text: str, modifications: List[TextModification]) -> str:
        """
        Applies a list of modifications to the text.
        
        CRITICAL LOGIC: 
        We must sort modifications in REVERSE order (descending start_pos).
        If we replace text at the beginning first, all subsequent indices 
        become invalid. Replacing from the end preserves earlier indices.
        """
        sorted_mods = sorted(modifications, key=lambda x: x.start_pos, reverse=True)
        modified_text = original_text
        for mod in sorted_mods:
            if mod.start_pos < 0 or mod.end_pos > len(modified_text):
                continue
            before = modified_text[:mod.start_pos]
            after = modified_text[mod.end_pos:]
            modified_text = before + mod.replacement_text + after
        return modified_text

    # ── TXT helper ────────────────────────────────────────────────────────────

    def generate_txt(self, text: str) -> io.BytesIO:
        """
        Returns the modified text as a clean UTF-8 .txt stream.
        Normalises whitespace: collapses runs of spaces within lines,
        preserves paragraph breaks.
        """
        clean = self._normalise_text(text)
        stream = io.BytesIO(clean.encode("utf-8"))
        stream.seek(0)
        return stream

    # ── DOCX helper ───────────────────────────────────────────────────────────

    def generate_docx(self, text: str) -> io.BytesIO:
        """
        Converts modified text to DOCX via a clean intermediate TXT
        representation.  This avoids the common artefacts caused by
        PDF/DOCX extraction (single words per line, stray newlines, etc.).

        Paragraph classification (applied after normalisation):
          ALL-CAPS line ≥ 5 chars               → Heading 1
          Starts with "N." or "N.N" + uppercase  → Heading 2
          Everything else                        → Normal body paragraph
        """
        paragraphs = self._to_paragraphs(text)

        doc = Document()
        # Remove the blank paragraph Word inserts by default
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)

        doc.add_heading("Modified Contract", 0)

        all_caps_re  = re.compile(r'^[A-Z][A-Z\s\-&/]{4,}$')
        numbered_re  = re.compile(r'^\d+(\.\d+)*\.\s+[A-Z]')

        for para in paragraphs:
            if not para:
                # Blank separator — small vertical space
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)

            elif all_caps_re.match(para):
                h = doc.add_heading(para, level=1)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after  = Pt(4)

            elif numbered_re.match(para):
                h = doc.add_heading(para, level=2)
                h.paragraph_format.space_before = Pt(8)
                h.paragraph_format.space_after  = Pt(2)

            else:
                p = doc.add_paragraph(para)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(4)

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

    # ── Private text normalisation ────────────────────────────────────────────

    @staticmethod
    def _normalise_text(text: str) -> str:
        """
        Produces a clean plain-text representation regardless of how the
        source was extracted (PDF word-per-line, DOCX, plain TXT).

        Strategy:
        1. Normalise line endings.
        2. Lines that look like headings (ALL-CAPS or numbered clause) are
           always kept as their own paragraph — insert a blank line around them.
        3. Soft-wrapped body lines (single newline between non-blank, non-heading
           text) are joined with a space.
        4. Collapse 3+ consecutive newlines to two (one blank line).
        5. Tidy internal whitespace within each line.
        """
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        all_caps_re = re.compile(r'^[A-Z][A-Z\s\-&/]{4,}$')
        numbered_re = re.compile(r'^\d+(\.\d+)*\.\s+[A-Z]')

        def is_heading(line: str) -> bool:
            s = line.strip()
            return bool(all_caps_re.match(s) or numbered_re.match(s))

        # Ensure headings are surrounded by blank lines so they survive
        # the soft-wrap join step as their own paragraph.
        lines = text.split('\n')
        guarded: List[str] = []
        for line in lines:
            if is_heading(line.strip()):
                if guarded and guarded[-1].strip() != '':
                    guarded.append('')
                guarded.append(line)
                guarded.append('')
            else:
                guarded.append(line)
        text = '\n'.join(guarded)

        # Join soft-wrapped lines: single \n between non-blank lines → space
        text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)

        # Collapse 3+ newlines to two
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Tidy each line: collapse internal whitespace, strip edges
        result_lines = [re.sub(r'[ \t]+', ' ', ln).strip() for ln in text.split('\n')]
        return '\n'.join(result_lines).strip()

    @classmethod
    def _to_paragraphs(cls, text: str) -> List[str]:
        """
        Normalise text and split into paragraph strings.
        Blank entries represent paragraph separators (preserved for spacing).
        """
        clean = cls._normalise_text(text)
        return clean.split('\n')

    # ── Diff summary ──────────────────────────────────────────────────────────

    def generate_diff_summary(self, modifications: List[TextModification]) -> dict:
        return {
            "total_changes": len(modifications),
            "sections_touched": [f"{m.start_pos}-{m.end_pos}" for m in modifications]
        }


    def generate_diff_summary(self, modifications: List[TextModification]) -> dict:
        """
        Optional: Returns a summary of what changed for audit logs.
        """
        return {
            "total_changes": len(modifications),
            "sections_touched": [f"{m.start_pos}-{m.end_pos}" for m in modifications]
        }